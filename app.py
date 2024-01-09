from flask import Flask, request, render_template, url_for, send_file, redirect, flash
from flask_login import LoginManager, UserMixin, login_user, logout_user, login_required
import json
import os
from docx import Document
from fpdf import FPDF
import subprocess
import re
import shutil
from tika import parser
from nltk.corpus import stopwords
import spacy
from spacy.matcher import Matcher
import bcrypt
import smtplib
from email.mime.text import MIMEText

app = Flask(__name__)

app.secret_key = 'PalmTree'
login_manager = LoginManager()
login_manager.init_app(app)

class User(UserMixin):
    def __init__(self, id):
        self.id = id

@login_manager.unauthorized_handler
def unauthorized_handler():
    return redirect(url_for('login'))

@login_manager.user_loader
def load_user(user_id):
    return User(user_id)

@app.route('/login', methods=['GET', 'POST'])
def login():
    if request.method == 'POST':
        username = request.form['username']
        password = request.form['password']
        
        with open('static/users', 'r') as file:
            users = json.load(file)
            if username in users:
                stored_password = users[username]['password']
                if bcrypt.checkpw(password.encode('utf-8'), stored_password.encode('utf-8')):
                    user = User(username)
                    login_user(user)
                    return redirect(url_for('index'))

        flash('Invalid username or password')
    
    return render_template('login.html')

@app.route('/logout')
def logout():
    logout_user()
    return redirect(url_for('login'))

json_file_path = 'static/uploads/parsed.json'
with open(json_file_path, 'r') as file:
    candidates_data = json.load(file)

def remove_duplicates(candidates_data):
    for candidate in candidates_data:
        for key, value in candidate.items():
            if isinstance(value, list):
                unique_values = []
                for item in value:
                    if item not in unique_values:
                        unique_values.append(item)
                candidate[key] = unique_values
    return candidates_data

candidates = remove_duplicates(candidates_data)

@app.route('/')
@login_required
def index():
    json_file_path = 'static/uploads/parsed.json'
    with open(json_file_path, 'r') as file:
        candidates_data = json.load(file)

    def remove_duplicates(candidates_data):
        for candidate in candidates_data:
            for key, value in candidate.items():
                if isinstance(value, list):
                    unique_values = []
                    for item in value:
                        if item not in unique_values:
                            unique_values.append(item)
                    candidate[key] = unique_values
        return candidates_data
    candidates = remove_duplicates(candidates_data)
    return render_template('index.html', candidates=candidates)

@app.route('/view',methods=['POST','GET'])
@login_required
def view():
    return render_template('view.html')

@app.route('/render_pdf/<id>')
@login_required
def render_pdf(id):
    pdf_path = f'static/uploads/after/{id}.pdf'
    return send_file(pdf_path)

@app.route('/mail', methods=['GET', 'POST'])
@login_required
def send_email():
    with open('static/search_results.json') as file:
        data = json.load(file)
        email_addresses = []
        for candidate in data:
            if 'E-Mail' in candidate:
                email_addresses.extend(candidate['E-Mail'])
        recipients = ', '.join(email_addresses)
    sender = "palmtreeresume@gmail.com"
    if request.method == 'POST':
        subject = request.form['subject']
        body = request.form['message']
        sender = "palmtreeresume@gmail.com"
        password = "qoix xmam ezkc bvau"
        done = 'Mail not sent'
        def send_email(subject, body, sender, recipient, password):
            msg = MIMEText(body)
            msg['Subject'] = subject
            msg['From'] = sender
            msg['To'] = recipient
            
            with smtplib.SMTP_SSL('smtp.gmail.com', 465) as smtp_server:
                smtp_server.login(sender, password)
                smtp_server.sendmail(sender, recipient, msg.as_string())

        for recipient in email_addresses:
            send_email(subject, body, sender, recipient, password)
        done='Mail sent!'
        return render_template('mail.html', done=done, to_email=recipients, from_email=sender)
    return render_template('mail.html', to_email=recipients, from_email=sender)

@app.route('/convert',methods=['POST','GET'])
@login_required
def convert_files():
    if request.method == 'POST':
        input_directory_path = 'static/uploads/before'
        output_directory_path = 'static/uploads/after'

        for filename in os.listdir(input_directory_path):
            input_file = os.path.join(input_directory_path, filename)
            output_file = os.path.join(output_directory_path, filename)
            if os.path.isdir(input_file):
                continue

            if input_file.lower().endswith('.pdf'):
                shutil.move(input_file, output_file)
            elif input_file.lower().endswith('.docx') or input_file.lower().endswith('.doc'):
                convert_to_pdf(input_file, output_file)    
        parse()
        return render_template('convert.html',done='Extracted all data')
    return render_template('convert.html')

def convert_to_pdf(input_file, output_file):
    # Read the input file
    try:
        if input_file.endswith('.docx'):
            document = Document(input_file)
            paragraphs = [paragraph.text for paragraph in document.paragraphs]
            content = '\n'.join(paragraphs)
        elif input_file.endswith('.doc'):
            try:
                # Convert .doc to .docx using unoconv
                docx_file = input_file + 'x'
                subprocess.run(['unoconv', '-f', 'docx', '-o', docx_file, input_file], capture_output=True)

                # Read the content from the converted .docx file
                doc = Document(docx_file)
                paragraphs = [paragraph.text for paragraph in doc.paragraphs]
                content = '\n'.join(paragraphs)

                # Delete the temporary .docx file
                os.remove(docx_file)
            except Exception as e:
                print(f"Error converting {input_file} to PDF: {str(e)}")
                return
        else:
            print("Invalid file format. Only .docx and .doc files are supported.")
            return

        # Create a PDF object and set the properties
        pdf = FPDF()
        pdf.add_page()
        pdf.set_auto_page_break(auto=True, margin=15)

        # Add the content to the PDF
        pdf.set_font('Arial', size=12)
        for line in content.splitlines():
            pdf.multi_cell(0, 10, line.encode('latin-1', 'ignore').decode('latin-1'))
        
        new_output = os.path.splitext(output_file)[0] + '.pdf'
        # Save the PDF to the output file
        pdf.output(new_output)
    except:
        return
    
@app.route('/upload', methods=['POST','GET'])
@login_required
def upload():
    return render_template('upload.html')

@app.route('/clear', methods=['POST'])
@login_required
def clear():
    # Clear the data
    with open("static/uploads/number.txt", "w") as f:
        f.write('0')
    folder_path = 'static/uploads/after'
    file_list = os.listdir(folder_path)
    for file_name in file_list:
        file_path = os.path.join(folder_path, file_name)
        if os.path.isfile(file_path):
            os.remove(file_path)
    folder_path = 'static/uploads/before'
    file_list = os.listdir(folder_path)
    for file_name in file_list:
        file_path = os.path.join(folder_path, file_name)
        if os.path.isfile(file_path):
            os.remove(file_path)
    filename = 'static/uploads/parsed.json'
    with open(filename, 'w') as file:
        file.write('[]')
    return render_template('index.html')

@app.route('/single_upload', methods=['GET', 'POST'])
@login_required
def single_upload():
    if request.method == 'POST':
        if 'resume' in request.files:
            file = request.files['resume']
            save_path = 'static/uploads/before'
            try:
                with open('static/uploads/number.txt', "r") as f:
                    number = int(f.read().strip())
                f.close()
            except FileNotFoundError:
                number = 0
            number += 1
            with open("static/uploads/number.txt", "w") as f:
                f.write(str(number))
            f.close()
            file_extension = os.path.splitext(file.filename)[1]
            output_file = os.path.join(save_path, str(number)+file_extension)
            file.save(output_file)
    return render_template('single_upload.html')

@app.route('/bulk_upload', methods=['POST', 'GET'])
@login_required
def bulk_upload():
    if request.method == 'POST':
        if 'files[]' in request.files:
            files = request.files.getlist('files[]')
            save_path = 'static/uploads/before'
            allowed_extensions = ['.doc', '.docx', '.pdf']

            try:
                with open('static/uploads/number.txt', "r") as f:
                    number = int(f.read().strip())
            except FileNotFoundError:
                number = 0

            for file in files:
                file_extension = os.path.splitext(file.filename)[1]
                if file_extension.lower() in allowed_extensions:
                    number += 1
                    output_file = os.path.join(save_path, str(number) + file_extension)
                    file.save(output_file)
                else:
                    return 'Invalid file extension. Only .doc, .docx, and .pdf files are allowed.'

            with open("static/uploads/number.txt", "w") as f:
                f.write(str(number))

            return 'Files saved successfully!'

        return 'No files found in the request.'

    return render_template('bulk_upload.html')


@app.route('/search', methods=['POST', 'GET'])
@login_required
def search():
    if request.method == 'POST':
        pdf_directory = 'static/uploads/after/'
        document_name = request.form['document_name']
        path = url_for('static', filename='uploads/after/{}.pdf'.format(document_name))
        return render_template('view.html', document=path)
    else:
        return render_template('view.html')
    
def parse():
    # Folder path containing the documents
    folder_path = 'static/uploads/after/'

    # List to store the extracted dictionaries
    parsed_dicts = []
    nlp = spacy.load('en_core_web_sm')
    STOPWORDS = set(stopwords.words('english'))


    # Iterate over the files in the folder
    for file_name in os.listdir(folder_path):
        file_path = os.path.join(folder_path, file_name)
        
        # Check if the file is a PDF
        if file_name.endswith('.pdf'):
            parsed = parser.from_file(file_path)
            text = parsed['content']
            if text is not None:
                text = text.replace('\n', ' ')
            extracted_text = {}

            # ID
            numeric_part = re.findall(r'\d+', file_name)
            if numeric_part:
                extracted_text['ID'] = numeric_part[0]
            
            #Name
            nlp = spacy.load('en_core_web_sm')

            matcher = Matcher(nlp.vocab)

            def extract_name(resume_text):
                if resume_text is None:
                    return None
                nlp_text = nlp(resume_text)
                pattern = [{'POS': 'PROPN'}]
                matcher.add('NAME', [pattern], on_match = None)
                matches = matcher(nlp_text)
                
                for match_id, start, end in matches:
                    span = nlp_text[start:end]
                    return span.text
            name = extract_name(text)
            if name:    
                extracted_text["Name"] = name
            
            #phone Number
            if text is not None and isinstance(text, str):
                phone_number = re.findall(r"\+?\d{1,3}[-.\s]?\(?\dGET{3}\)?[-.\s]?\d{3}[-.\s]?\d{4}", text)
                if phone_number:
                    extracted_text["Phone Number"] = phone_number

            #email
            if text is not None and isinstance(text, str):
                email = re.findall(r"\b[A-Za-z0-9._%+-]+@[A-Za-z0-9.-]+\.[A-Za-z]{2,}\b", text)
                if email:
                    extracted_text["E-Mail"] = email

            #education
            EDUCATION = [
                        'BE','B.E.', 'B.E', 'BS', 'B.S', 'B.S.'
                        'ME', 'M.E', 'M.E.', 'MS', 'M.S', 'M.S.'
                        'BTECH', 'B.TECH', 'M.TECH', 'MTECH', 
                        'SSC', 'HSC', 'CBSE', 'ICSE', 'X', 'XII'
                    ]
            if text is not None and isinstance(text, str):
                def extract_education(resume_text):
                    nlp_text = nlp(resume_text)
                    nlp_text = [sent.text.strip() for sent in nlp_text.sents]
                    edu = {}
                    for index, text in enumerate(nlp_text):
                        for tex in text.split():
                            tex = re.sub(r'[?|$|.|!|,]', r'', tex)
                            if tex.upper() in EDUCATION and tex not in STOPWORDS:
                                if index + 1 < len(nlp_text):
                                    edu[tex] = text + nlp_text[index + 1]
                    education = []
                    for key in edu.keys():
                        year = re.search(re.compile(r'(((20|19)(\d{2})))'), edu[key])
                        if year:
                            education.append((key, ''.join(year[0])))
                        else:
                            education.append(key)
                    return education
                education = extract_education(text)
                if education:
                    extracted_text["Qualification"] = education
                
            #university
            sub_patterns = ['[A-Z][a-z]* University','[A-Z][a-z]* Educational Institute',
                            'University of [A-Z][a-z]*',
                            '[A-Z][a-z]* College']
            if text is not None and isinstance(text, str):
                pattern = '({})'.format('|'.join(sub_patterns))
                matches = re.findall(pattern, text)
                if matches:
                    extracted_text["Institutes"] = matches

            #company
            if text is not None and isinstance(text, str):
                pattern = r"\b([A-Za-z\s]+?)\s?(?:Pvt Ltd|LLC|Inc|International|Corporation)\b"
                Exp = re.findall(pattern, text)
                if Exp:
                    extracted_text["Company"] = Exp

            #skills
            if text is not None and isinstance(text, str):
                def extract_skills(resume_text):
                    nlp_text = nlp(resume_text)
                    
                    file_path = 'static/skills.txt'
                    skills = []
                    with open(file_path, 'r') as file:
                        for line in file:
                            skills.append(line.strip()) 
                    return skills
                skills = extract_skills(text)
                skills
                all_skills = skills
                pattern = r'\b(?:{})\b'.format('|'.join(map(re.escape, skills)))
                skill = re.findall(pattern, text,flags=re.IGNORECASE)
                skill = list(set(skill))
                if skills:
                    extracted_text["Skills"] = skill

            
            # Append the extracted dictionary to the list
            parsed_dicts.append(extracted_text)

    # Save the parsed_dicts list to a file
    filename = 'static/uploads/parsed.json'
    with open(filename, 'w') as file:
        json.dump(parsed_dicts, file)


def search_by_skills(conditions, dictionary_list):
    results = []
    for dictionary in dictionary_list:
        if "Skills" in dictionary:
            skills = set(dictionary["Skills"])
            if evaluate_conditions(conditions, skills):
                results.append(dictionary)
    return results

def evaluate_conditions(conditions, skills):
    operator = "&"
    condition_groups = []
    current_group = []
    for condition in conditions:
        if condition == "&":
            operator = "&"
        elif condition == "|":
            operator = "|"
        else:
            current_group.append(condition)
        if operator == "|":
            condition_groups.append(current_group)
            current_group = []
    if current_group:
        condition_groups.append(current_group)

    return any(evaluate_group(group, skills, operator) for group in condition_groups)

def evaluate_group(group, skills, operator):
    if operator == "&":
        return all(re.search(re.escape(condition), " ".join(skills), re.IGNORECASE) for condition in group)
    elif operator == "|":
        return any(re.search(re.escape(condition), " ".join(skills), re.IGNORECASE) for condition in group)
    return False

@app.route('/searchresume', methods=['GET', 'POST'])
@login_required
def search_results():
    json_file_path = 'static/uploads/parsed.json'
    with open(json_file_path, 'r') as file:
        candidates_data = json.load(file)

    def remove_duplicates(candidates_data):
        for candidate in candidates_data:
            for key, value in candidate.items():
                if isinstance(value, list):
                    unique_values = []
                    for item in value:
                        if item not in unique_values:
                            unique_values.append(item)
                    candidate[key] = unique_values
        return candidates_data
    #candidates = remove_duplicates(candidates_data)
    if request.method == 'POST':
        search_query = request.form.get('search_query')
        conditions = search_query.split()
        results = search_by_skills(conditions, candidates_data)
        with open('static/search_results.json', 'w') as file:
            json.dump(results, file)
        return render_template('searchresume.html', results=results)
    return render_template('searchresume.html')

if __name__ == '__main__':
    app.run(debug=True,host='0.0.0.0')



